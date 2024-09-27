from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, send_file
from models import db, User, VocabularyEntry, ContextEntry, Exercise
from extensions import bcrypt
from nlp_utils import (find_target_word, lemmatize_sentence, check_linguistic_features, clean_text,
                       generate_and_validate_contexts, generate_contexts_no_validation, strip_prefixes, update_pos)

from sqlalchemy.sql import func
from flask_migrate import Migrate
from flask_wtf.csrf import CSRFProtect, CSRFError
import pandas as pd
import os
import re
from datetime import datetime
import nltk
from nltk import pos_tag
from nltk.tokenize import word_tokenize
from nltk.corpus import wordnet
from sqlalchemy import or_
import logging
import json

from werkzeug.utils import secure_filename
from flask import request
import magic  # libmagic wrapper to check MIME types of uploaded files

from google.api_core.exceptions import ResourceExhausted #note if there is a problem with Gemini quota and use Mistral instead
import openai #for Mistral generation

# in order to export files to Word
from docx import Document
import io
from io import BytesIO

from datetime import timedelta

from functools import wraps #to use decorator to check if user is logged in


nltk.download('averaged_perceptron_tagger')
nltk.download('punkt')
nltk.download('wordnet')

app = Flask(__name__)
csrf = CSRFProtect(app)
app.config['SECRET_KEY'] = 'fj90if0'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///site.db'
app.config['UPLOAD_FOLDER'] = 'uploads'  # Folder to store uploaded files
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(minutes=1440) #make sure user is logged out automatically after a number of minutes of inactivity; 24 hours currently set 


# Initialize extensions with the Flask app
db.init_app(app)
migrate = Migrate(app, db)
#bcrypt = Bcrypt(app)
bcrypt.init_app(app)
logging.basicConfig(level=logging.DEBUG)


import time
import spacy
import nltk
import google.generativeai as genai
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.stem import PorterStemmer
from nltk.corpus import wordnet
from flask import jsonify
import langid
import random
from difflib import SequenceMatcher
from nltk.stem import WordNetLemmatizer
import inflect
nltk.download('punkt')
nltk.download('wordnet')

lemmatizer = WordNetLemmatizer()

nlp = spacy.load("en_core_web_sm")

# Initialize the inflect engine - to get inflected forms of distractors if needed
p = inflect.engine()

#genai.configure(api_key="AIzaSyAFsbEx5fpzsoePD_1Wyct63A8PgBejBKI")  
#model = genai.GenerativeModel('gemini-pro')


# Configuring maximum upload size (e.g., 10MB)
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024  # 10 Megabytes

ALLOWED_EXTENSIONS = {'csv', 'xlsx'}
ALLOWED_MIME_TYPES = {'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'application/vnd.ms-excel', 'text/csv'}


def login_required(f):  #to ensure user is logged in
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'username' not in session:
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def allowed_mime_type(file):
    mime = magic.Magic(mime=True)
    file_mime_type = mime.from_buffer(file.read(1024))  # Check the first 1024 bytes
    file.seek(0)  # Reset the file pointer to the start
    return file_mime_type in ALLOWED_MIME_TYPES


@app.route('/')
def home():
    return render_template('home.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        email = request.form['email']
        password = request.form['password']

        existing_user = User.query.filter((User.username == username) | (User.email == email)).first()
        if existing_user:
            flash("Username or email already registered.", 'error')
            return render_template('register.html')

        new_user = User(username=username, email=email)
        new_user.set_password(password)
        db.session.add(new_user)
        db.session.commit()
        flash('Registration successful. Please log in.', 'success')
        return redirect(url_for('login'))

    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if 'username' in session:
        return redirect(url_for('home'))

    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')

        user = User.query.filter_by(username=username).first()
        if user and user.check_password(password):
            session.permanent = True  # Make the session permanent
            session['username'] = username
            session['user_id'] = user.id
            return redirect(url_for('home'))
        else:
            flash('Invalid username or password', 'error')

    return render_template('login.html')



@app.route('/logout', methods=['GET', 'POST'])
def logout():
    session.clear()
    return redirect(url_for('logout_confirmation'))

@app.route('/logout_confirmation')
def logout_confirmation():
    return """
    <p>You have been logged out successfully.</p>
    <p><a href="/">Back to Home Page</a></p>
    """
    
@app.route('/personal_space/<username>')
def personal_space(username):
    if 'username' not in session or session['username'] != username:
        return redirect(url_for('login'))

    redirect_back = request.args.get('redirect_back')
    # Log the redirect_back value
    app.logger.info(f"Redirect back URL in personal_space: {redirect_back}")

    return render_template('personal_space.html', username=username, redirect_back=redirect_back)




@app.route('/about')
def about():
    return render_template('about.html')


@app.route('/manage_vocabulary_lists', methods=['GET', 'POST'])
@login_required
@csrf.exempt
def manage_vocabulary_lists():
    if request.method == 'POST':
        # Reset the 'used_as_distractor' column for all entries
        VocabularyEntry.query.update({'used_as_distractor': 'no'})
        db.session.commit()
        flash('The list of utilized distractors has been reset.', 'success')
        return redirect(url_for('manage_vocabulary_lists'))

    return render_template('manage_vocabulary_lists.html')



@app.route('/add_vocabulary_list', methods=['GET', 'POST'])
@login_required
def add_vocabulary_list():
    if request.method == 'POST':
        course_code = request.form['course_code']
        cefr_level = request.form['cefr_level']
        domain = request.form['domain']
        file = request.files['file']

        if file:
            app.logger.info(f'Received POST request with file: {file.filename}')
            if not allowed_file(file.filename):
                flash('Invalid file format. Please upload a .csv or .xlsx file.', 'error')
                return redirect(url_for('add_vocabulary_list'))
            
            if not allowed_mime_type(file):
                flash('Invalid file MIME type. Please upload a file with the correct format.', 'error')
                return redirect(url_for('add_vocabulary_list'))
            
            filename = secure_filename(file.filename)
            file_content = file.read()

            try:
                if filename.endswith('.xlsx'):
                    df = pd.read_excel(io.BytesIO(file_content))
                    app.logger.info('File processed as .xlsx')
                elif filename.endswith('.csv'):
                    df = pd.read_csv(io.StringIO(file_content.decode('utf-8')))
                    app.logger.info('File processed as .csv')
                else:
                    flash('Invalid file format. Please upload a .csv or .xlsx file.', 'error')
                    return redirect(url_for('add_vocabulary_list'))
            except Exception as e:
                flash(f'Error reading the file: {str(e)}', 'error')
                app.logger.error(f'Error reading file: {e}')
                return redirect(url_for('add_vocabulary_list'))

            # Check for 'Selected' and 'selected' columns and handle accordingly
            if 'Selected' in df.columns:
                df.rename(columns={'Selected': 'selected'}, inplace=True)
            if 'selected' not in df.columns:
                df['selected'] = 'no'
                
            # Ensure 'used_as_distractor' column exists
            df['used_as_distractor'] = 'no'

            required_columns = ['Item', 'POS', 'Translation', 'Lesson title', 'Reading or listening']
            for column in required_columns:
                if column not in df.columns:
                    df[column] = ''

            user_id = session['user_id']  # Use user_id from the session
            date_loaded = datetime.utcnow()
            number_of_contexts = 0

            app.logger.info('Starting database insertion.')

            for index, row in df.iterrows():
                try:
                    entry = VocabularyEntry(
                        item=strip_prefixes(row['Item']),
                        pos=row['POS'] if row['POS'] else None,
                        translation=row['Translation'],
                        lesson_title=row['Lesson title'],
                        reading_or_listening=row['Reading or listening'],
                        course_code=course_code,
                        cefr_level=cefr_level,
                        domain=domain,
                        user_id=user_id,  # Replace user with user_id
                        date_loaded=date_loaded,
                        number_of_contexts=number_of_contexts,
                        selected=row['selected'],  # Retain value from the file or 'no' if not present
                        used_as_distractor='no'  # Default to 'no'
                    )
                    db.session.add(entry)
                except Exception as e:
                    flash(f'Error inserting row {index + 1}: {str(e)}', 'error')
                    app.logger.error(f'Error inserting row {index + 1}: {e}')

            try:
                db.session.commit()
                flash('Thank you for uploading a vocabulary list!', 'success')
                app.logger.info('Database insertion completed successfully.')
            except Exception as e:
                db.session.rollback()
                flash(f'Error committing to the database: {str(e)}', 'error')
                app.logger.error(f'Error committing to the database: {e}')

            return redirect(url_for('home'))
        else:
            flash('No file provided.', 'error')
            return redirect(url_for('add_vocabulary_list'))

    return render_template('add_vocabulary_list.html')


@app.errorhandler(413)
def request_entity_too_large(error):
    flash('The file is too large. Please upload a file smaller than 10MB.', 'error')
    return redirect(url_for('add_vocabulary_list'))


def allowed_mime_type(file):
    ALLOWED_MIME_TYPES = ['text/csv', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']
    return file.mimetype in ALLOWED_MIME_TYPES


@app.route('/manage_current_entries', methods=['GET', 'POST'])
@csrf.exempt
@login_required
def manage_current_entries():
    # Clear session parameters on a fresh load if not coming from 'edit_vocabulary_entry' or other internal navigation
    if request.method == 'GET' and not request.referrer:
        session.pop('selected_params', None)
    elif request.method == 'GET' and 'edit_vocabulary_entry' not in request.referrer:
        session.pop('selected_params', None)

    # Get unique values for filters
    course_codes = sorted([code.course_code for code in VocabularyEntry.query.with_entities(VocabularyEntry.course_code).distinct().all()])
    domains = sorted([domain.domain for domain in VocabularyEntry.query.with_entities(VocabularyEntry.domain).distinct().all()])
    levels = sorted([level.cefr_level for level in VocabularyEntry.query.with_entities(VocabularyEntry.cefr_level).distinct().all()])
    context_ranges = ['less_than_10', '10_to_20', '20_to_50', '50_to_100']
    items = sorted([strip_prefixes(item.item) for item in VocabularyEntry.query.with_entities(VocabularyEntry.item).distinct().all()])

    results = []
    selected_params = session.get('selected_params', {})

    # If the form is submitted, process the form data and save to session
    if request.method == 'POST':
        search_params = {
            'course_code': request.form.getlist('course_code'),
            'domain': request.form.getlist('domain'),
            'cefr_level': request.form.getlist('cefr_level'),
            'item': request.form.get('item'),
            'number_of_contexts': request.form.getlist('number_of_contexts'),
            'selected': request.form.getlist('selected')
        }
        selected_params = search_params
        session['selected_params'] = selected_params  # Store in session for back navigation

    # Rebuild the query using selected parameters
    if selected_params:
        query = VocabularyEntry.query
        if selected_params.get('course_code'):
            query = query.filter(VocabularyEntry.course_code.in_(selected_params['course_code']))
        if selected_params.get('domain'):
            query = query.filter(VocabularyEntry.domain.in_(selected_params['domain']))
        if selected_params.get('cefr_level'):
            query = query.filter(VocabularyEntry.cefr_level.in_(selected_params['cefr_level']))
        if selected_params.get('item'):
            query = query.filter(VocabularyEntry.item == selected_params['item'])
        if selected_params.get('number_of_contexts'):
            context_filters = []
            for context_range in selected_params['number_of_contexts']:
                if context_range == 'less_than_10':
                    context_filters.append(VocabularyEntry.number_of_contexts < 10)
                elif context_range == '10_to_20':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(10, 20))
                elif context_range == '20_to_50':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(20, 50))
                elif context_range == '50_to_100':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(50, 100))
            query = query.filter(or_(*context_filters))
        if selected_params.get('selected'):
            query = query.filter(VocabularyEntry.selected.in_(selected_params['selected']))

        results = query.order_by(VocabularyEntry.item).all()

    return render_template('manage_current_entries.html',
                           course_codes=course_codes,
                           domains=domains,
                           levels=levels,
                           context_ranges=context_ranges,
                           items=items,
                           results=results,
                           selected_params=selected_params)







@app.route('/toggle_selection/<int:item_id>', methods=['POST'])
@login_required
@csrf.exempt
def toggle_selection(item_id):
    selected = request.form.get('selected')
    item = VocabularyEntry.query.get_or_404(item_id)
    item.selected = selected
    db.session.commit()
    return redirect(url_for('manage_current_entries'))


@app.route('/delete_vocabulary_entry/<int:item_id>', methods=['POST'])
@login_required
@csrf.exempt
def delete_vocabulary_entry(item_id):
    item = VocabularyEntry.query.get_or_404(item_id)
    db.session.delete(item)
    db.session.commit()
    flash(f'The item "{item_name}" has been deleted from the database.', 'success')
    return redirect(url_for('manage_current_entries'))




@app.route('/edit_vocabulary_entry/<int:item_id>', methods=['GET', 'POST'])
@login_required
@csrf.exempt
def edit_vocabulary_entry(item_id):
    item = VocabularyEntry.query.get_or_404(item_id)

    # Get selected_params from the request args (or defaults)
    selected_params = {
        'course_code': request.args.getlist('course_code'),
        'domain': request.args.getlist('domain'),
        'cefr_level': request.args.getlist('cefr_level'),
        'item': request.args.get('item'),
        'number_of_contexts': request.args.getlist('number_of_contexts'),
        'selected': request.args.getlist('selected')
    }

    # Filter out any empty parameters
    selected_params = {k: v for k, v in selected_params.items() if v}

    if request.method == 'POST':
        item.item = request.form.get('item')
        item.pos = request.form.get('pos')
        item.translation = request.form.get('translation')
        item.course_code = request.form.get('course_code')
        item.cefr_level = request.form.get('cefr_level')
        item.domain = request.form.get('domain')
        item.lesson_title = request.form.get('lesson_title')
        item.reading_or_listening = request.form.get('reading_or_listening')

        db.session.commit()
        return redirect(url_for('manage_current_entries', **selected_params))

    return render_template('edit_vocabulary_entry.html', item=item, selected_params=selected_params)









@app.route('/update_vocabulary_entry/<int:item_id>', methods=['POST'])
@login_required
def update_vocabulary_entry(item_id):
    item = VocabularyEntry.query.get_or_404(item_id)
    item.item = request.form.get('item')
    item.pos = request.form.get('pos')
    item.translation = request.form.get('translation')
    item.course_code = request.form.get('course_code')
    item.cefr_level = request.form.get('cefr_level')
    item.domain = request.form.get('domain')
    item.lesson_title = request.form.get('lesson_title')
    item.reading_or_listening = request.form.get('reading_or_listening')

    db.session.commit()
    return redirect(url_for('manage_current_entries'))





@app.route('/add_context_examples', methods=['GET', 'POST'])
@csrf.exempt
@login_required
def add_context_examples():
    # Get unique values for dropdowns and sort them alphabetically
    course_codes = sorted([code.course_code for code in VocabularyEntry.query.with_entities(VocabularyEntry.course_code).distinct().all()])
    domains = sorted([domain.domain for domain in VocabularyEntry.query.with_entities(VocabularyEntry.domain).distinct().all()])
    levels = sorted([level.cefr_level for level in VocabularyEntry.query.with_entities(VocabularyEntry.cefr_level).distinct().all()])
    items = sorted([strip_prefixes(item.item) for item in VocabularyEntry.query.with_entities(VocabularyEntry.item).distinct().all()])

    # Initialize results and search parameters
    results = []
    selected_params = {
        'course_code': [],
        'domain': [],
        'cefr_level': [],
        'number_of_contexts': [],
        'item': '',
        'selected': []  # Added selected filter
    }

    # Check for GET parameters if coming back from generate_context
    if request.method == 'GET':
        selected_params['course_code'] = request.args.getlist('course_code')
        selected_params['domain'] = request.args.getlist('domain')
        selected_params['cefr_level'] = request.args.getlist('cefr_level')
        selected_params['number_of_contexts'] = request.args.getlist('number_of_contexts')
        selected_params['item'] = request.args.get('item', '')
        selected_params['selected'] = request.args.getlist('selected')

        # Automatically execute the search based on the query parameters
        if any(selected_params.values()):
            query = VocabularyEntry.query

            if selected_params['course_code']:
                query = query.filter(VocabularyEntry.course_code.in_(selected_params['course_code']))
            if selected_params['domain']:
                query = query.filter(VocabularyEntry.domain.in_(selected_params['domain']))
            if selected_params['cefr_level']:
                query = query.filter(VocabularyEntry.cefr_level.in_(selected_params['cefr_level']))
            if selected_params['item']:
                query = query.filter(VocabularyEntry.item == selected_params['item'])
            if selected_params['number_of_contexts']:
                context_filters = []
                for context_range in selected_params['number_of_contexts']:
                    if context_range == 'less_than_10':
                        context_filters.append(VocabularyEntry.number_of_contexts < 10)
                    elif context_range == '10_to_20':
                        context_filters.append(VocabularyEntry.number_of_contexts.between(10, 20))
                    elif context_range == '20_to_50':
                        context_filters.append(VocabularyEntry.number_of_contexts.between(20, 50))
                    elif context_range == '50_to_100':
                        context_filters.append(VocabularyEntry.number_of_contexts.between(50, 100))
                query = query.filter(or_(*context_filters))
            if selected_params['selected']:
                query = query.filter(VocabularyEntry.selected.in_(selected_params['selected']))

            results = query.order_by(VocabularyEntry.item).all()

            # Add actual count of context examples to each result
            for result in results:
                context_count = ContextEntry.query.filter_by(item_id=result.id).count()
                result.number_of_contexts = context_count

    # Handle POST search requests
    if request.method == 'POST':
        search_params = {
            'course_code': request.form.getlist('course_code'),
            'domain': request.form.getlist('domain'),
            'cefr_level': request.form.getlist('cefr_level'),
            'item': request.form.get('item'),
            'number_of_contexts': request.form.getlist('number_of_contexts'),
            'selected': request.form.getlist('selected')  # Get selected values
        }
        selected_params = search_params

        query = VocabularyEntry.query
        if search_params['course_code']:
            query = query.filter(VocabularyEntry.course_code.in_(search_params['course_code']))
        if search_params['domain']:
            query = query.filter(VocabularyEntry.domain.in_(search_params['domain']))
        if search_params['cefr_level']:
            query = query.filter(VocabularyEntry.cefr_level.in_(search_params['cefr_level']))
        if search_params['item']:
            query = query.filter(VocabularyEntry.item == search_params['item'])
        if search_params['number_of_contexts']:
            context_filters = []
            for context_range in search_params['number_of_contexts']:
                if context_range == 'less_than_10':
                    context_filters.append(VocabularyEntry.number_of_contexts < 10)
                elif context_range == '10_to_20':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(10, 20))
                elif context_range == '20_to_50':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(20, 50))
                elif context_range == '50_to_100':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(50, 100))
            query = query.filter(or_(*context_filters))
        if search_params['selected']:
            query = query.filter(VocabularyEntry.selected.in_(search_params['selected']))

        results = query.order_by(VocabularyEntry.item).all()

        # Add actual count of context examples to each result
        for result in results:
            context_count = ContextEntry.query.filter_by(item_id=result.id).count()
            result.number_of_contexts = context_count

    return render_template('add_context_examples.html',
                           course_codes=course_codes,
                           domains=domains,
                           levels=levels,
                           items=items,
                           results=results,
                           selected_params=selected_params)







@app.route('/current_contexts/<int:item_id>', methods=['GET'])
@login_required
def current_contexts(item_id):
    item = VocabularyEntry.query.get_or_404(item_id)
    contexts = ContextEntry.query.filter_by(item_id=item_id).all()
    return render_template('current_contexts.html', item=item, contexts=contexts)



@app.route('/generate_context/<int:item_id>', methods=['GET', 'POST'])
@csrf.exempt
@login_required
def generate_context(item_id):
    item = VocabularyEntry.query.get_or_404(item_id)
    generated_contexts = None

    # Get the previous search parameters from query string
    course_code = request.args.get('course_code')
    domain = request.args.get('domain')
    cefr_level = request.args.get('cefr_level')
    number_of_contexts = request.args.get('number_of_contexts')
    item_search = request.args.get('item')
    selected = request.args.get('selected')

    if item.number_of_contexts >= 100:
        generate_disabled = True
    else:
        generate_disabled = False

    if request.method == 'POST' and not generate_disabled:
        if 'generate_contexts' in request.form:
            pos = item.pos
            domain = item.domain
            level = item.cefr_level
            generated_contexts = generate_and_validate_contexts(item.item, pos, domain, level)
        
        elif 'generate_contexts_no_checks' in request.form:  # New button action
            pos = item.pos
            domain = item.domain
            level = item.cefr_level
            generated_contexts = generate_contexts_no_validation(item.item, pos, domain, level)

        elif 'save_contexts' in request.form:
            selected_contexts = request.form.getlist('selected_contexts')
            context_texts = request.form.getlist('context_texts')
            user_id = session.get('user_id')  # Use user_id from the session
            date_added = datetime.utcnow()
            generation_methods = request.form.getlist('generation_methods')

            # Check if the new contexts would push the total over 100
            total_contexts_after_save = item.number_of_contexts + len(selected_contexts)
            if total_contexts_after_save > 100:
                flash('Cannot save contexts. The total number of contexts would exceed the limit of 100.', 'danger')
                return redirect(url_for('generate_context', item_id=item_id))

            for idx in selected_contexts:
                idx = int(idx) - 1
                if idx < len(context_texts):
                    context = context_texts[idx].strip()
                    target_word = find_target_word(context, item.item)

                    # Skip the context if no valid target word was found
                    if target_word is None:
                        app.logger.warning(f"Skipping context: {context} - No valid target word found.")
                        continue

                    # Ensure capitalization of target_word if it's the first word
                    if context.lower().startswith(target_word.lower()):
                        context = target_word.capitalize() + context[len(target_word):]

                    generation_method = generation_methods[idx]
                    new_context = ContextEntry(
                        item_id=item_id,
                        context=context,
                        target_word=target_word,
                        user_id=user_id,  # Replace user with user_id
                        date_added=date_added,
                        generation_method=generation_method
                    )
                    db.session.add(new_context)

            # Update the number_of_contexts field in VocabularyEntry
            item.number_of_contexts = ContextEntry.query.filter_by(item_id=item_id).count()
            db.session.commit()
            flash('Selected contexts saved successfully.', 'success')

            # Check again if the context limit has been reached after saving
            if item.number_of_contexts >= 100:
                generate_disabled = True

    # Get current contexts with pagination
    page = request.args.get('page', 1, type=int)
    per_page = 10
    pagination = ContextEntry.query.filter_by(item_id=item_id).paginate(page=page, per_page=per_page, error_out=False)
    current_contexts = pagination.items

    # Find and highlight target words with correct capitalization
    highlighted_contexts = []
    for context in current_contexts:
        target_word = find_target_word(context.context, item.item)

        # Skip the context if no valid target word was found
        if target_word is None:
            app.logger.warning(f"Skipping context: {context.context} - No valid target word found.")
            continue

        # Check if target word is at the start and capitalize it if necessary
        if context.context.lower().startswith(target_word.lower()):
            context.context = target_word.capitalize() + context.context[len(target_word):]

        # Highlight the target word while preserving its original case
        highlighted_context = re.sub(
            f"(?i)\\b{target_word}\\b",  # (?i) makes the regex case-insensitive but preserves original case
            lambda match: f"<b>{match.group(0)}</b>",  # Use a lambda to preserve the case of the matched word
            context.context
        )

        highlighted_contexts.append(highlighted_context)

    return render_template('generate_context.html', item=item, current_contexts=highlighted_contexts, generated_contexts=generated_contexts, pagination=pagination, per_page=per_page, page=page, generate_disabled=generate_disabled, find_target_word=find_target_word, course_code=course_code, domain=domain, cefr_level=cefr_level, number_of_contexts=number_of_contexts, item_search=item_search, selected=selected)




@app.route('/create_exercises', methods=['GET', 'POST'])
@csrf.exempt
@login_required
def create_exercises():
    # Get unique values for checkboxes and sort them alphabetically
    course_codes = sorted([code.course_code for code in VocabularyEntry.query.with_entities(VocabularyEntry.course_code).distinct().all()])
    domains = sorted([domain.domain for domain in VocabularyEntry.query.with_entities(VocabularyEntry.domain).distinct().all()])
    levels = sorted([level.cefr_level for level in VocabularyEntry.query.with_entities(VocabularyEntry.cefr_level).distinct().all()])
    context_ranges = ['less_than_10', '10_to_20', '20_to_50', '50_to_100']

    items = sorted([strip_prefixes(item.item) for item in VocabularyEntry.query.with_entities(VocabularyEntry.item).distinct().all()])

    results = []  # Initially, no results
    selected_params = {}

    # Handle POST request (when the user submits the search form)
    if request.method == 'POST':
        search_params = {
            'course_code': request.form.getlist('course_code'),
            'domain': request.form.getlist('domain'),
            'cefr_level': request.form.getlist('cefr_level'),
            'number_of_contexts': request.form.getlist('number_of_contexts'),
            'selected': request.form.getlist('selected'),
            'item': request.form.get('item')  # Add item to search parameters
        }
        selected_params = search_params

        query = VocabularyEntry.query
        if search_params['course_code']:
            query = query.filter(VocabularyEntry.course_code.in_(search_params['course_code']))
        if search_params['domain']:
            query = query.filter(VocabularyEntry.domain.in_(search_params['domain']))
        if search_params['cefr_level']:
            query = query.filter(VocabularyEntry.cefr_level.in_(search_params['cefr_level']))
        if search_params['number_of_contexts']:
            context_filters = []
            for context_range in search_params['number_of_contexts']:
                if context_range == 'less_than_10':
                    context_filters.append(VocabularyEntry.number_of_contexts < 10)
                elif context_range == '10_to_20':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(10, 20))
                elif context_range == '20_to_50':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(20, 50))
                elif context_range == '50_to_100':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(50, 100))
            query = query.filter(or_(*context_filters))
        if search_params['selected']:
            query = query.filter(VocabularyEntry.selected.in_(search_params['selected']))
        if search_params['item']:  # Add this condition to filter by item
            query = query.filter(VocabularyEntry.item == search_params['item'])  # Exact match on the item name

        results = query.order_by(VocabularyEntry.item).all()  # Get results based on search

        # Redirect to GET request to update the URL with search params
        return redirect(url_for('create_exercises',
                                course_code=search_params['course_code'],
                                domain=search_params['domain'],
                                cefr_level=search_params['cefr_level'],
                                number_of_contexts=search_params['number_of_contexts'],
                                selected=search_params['selected'],
                                item=search_params['item']))  # Add item to redirect parameters

    # Handle GET request (when the user clicks "Back" or accesses via URL)
    elif request.method == 'GET':
        selected_params = {
            'course_code': request.args.getlist('course_code'),
            'domain': request.args.getlist('domain'),
            'cefr_level': request.args.getlist('cefr_level'),
            'number_of_contexts': request.args.getlist('number_of_contexts'),
            'selected': request.args.getlist('selected'),
            'item': request.args.get('item')  # Add item to GET parameters
        }

        if any(selected_params.values()):  # Only run the query if there are any filters selected
            query = VocabularyEntry.query
            if selected_params['course_code']:
                query = query.filter(VocabularyEntry.course_code.in_(selected_params['course_code']))
            if selected_params['domain']:
                query = query.filter(VocabularyEntry.domain.in_(selected_params['domain']))
            if selected_params['cefr_level']:
                query = query.filter(VocabularyEntry.cefr_level.in_(selected_params['cefr_level']))
            if selected_params['number_of_contexts']:
                context_filters = []
                for context_range in selected_params['number_of_contexts']:
                    if context_range == 'less_than_10':
                        context_filters.append(VocabularyEntry.number_of_contexts < 10)
                    elif context_range == '10_to_20':
                        context_filters.append(VocabularyEntry.number_of_contexts.between(10, 20))
                    elif context_range == '20_to_50':
                        context_filters.append(VocabularyEntry.number_of_contexts.between(20, 50))
                    elif context_range == '50_to_100':
                        context_filters.append(VocabularyEntry.number_of_contexts.between(50, 100))
                query = query.filter(or_(*context_filters))
            if selected_params['selected']:
                query = query.filter(VocabularyEntry.selected.in_(selected_params['selected']))
            if selected_params['item']:  # Add this condition to filter by item
                query = query.filter(VocabularyEntry.item == selected_params['item'])  # Exact match on the item name

            results = query.order_by(VocabularyEntry.item).all()  # Get filtered results

    return render_template('create_exercises.html',
                           course_codes=course_codes,
                           domains=domains,
                           levels=levels,
                           context_ranges=context_ranges,
                           items=items,
                           results=results,
                           selected_params=selected_params)






@app.route('/filter_items', methods=['POST'])
def filter_items():
    try:
        data = request.get_json()
        logging.debug(f"Received data: {data}")
        course_codes = data.get('course_code', [])
        domains = data.get('domain', [])
        cefr_levels = data.get('cefr_level', [])
        context_ranges = data.get('number_of_contexts', [])

        query = VocabularyEntry.query

        if course_codes:
            query = query.filter(VocabularyEntry.course_code.in_(course_codes))
        if domains:
            query = query.filter(VocabularyEntry.domain.in_(domains))
        if cefr_levels:
            query = query.filter(VocabularyEntry.cefr_level.in_(cefr_levels))
        if context_ranges:
            context_filters = []
            for context_range in context_ranges:
                if context_range == 'less_than_10':
                    context_filters.append(VocabularyEntry.number_of_contexts < 10)
                elif context_range == '10_to_20':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(10, 20))
                elif context_range == '20_to_50':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(20, 50))
                elif context_range == '50_to_100':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(50, 100))
            query = query.filter(or_(*context_filters))

        items = sorted([strip_prefixes(item.item) for item in query.with_entities(VocabularyEntry.item).distinct().all()])

        logging.debug(f"Filtered items: {items}")
        return jsonify(items)
    except Exception as e:
        logging.error(f"Error in filter_items: {e}")
        return jsonify({'error': str(e)}), 400



@app.route('/create_gapfill/<int:context_id>', methods=['GET', 'POST'])
@csrf.exempt  # Disable CSRF protection for this route
@login_required
def create_gapfill(context_id):
    # Retrieve the context and related vocabulary entry
    context = ContextEntry.query.get_or_404(context_id)
    item = VocabularyEntry.query.get_or_404(context.item_id)
    
    # Retrieve the 'redirect_back' parameter, defaulting to the 'generate_exercises' page if not provided
    redirect_back = request.args.get('redirect_back', url_for('generate_exercises', item_id=item.id))

    # Log various details for debugging
    app.logger.debug(f"Creating gapfill for context ID: {context_id}")
    app.logger.debug(f"Associated vocabulary item: {item.item}")
    app.logger.debug(f"Context: {context.context}")
    app.logger.debug(f"Redirect back URL: {redirect_back}")
    
    target_word = context.target_word
    translation = item.translation

    # Process the context to handle cases where 'a' or 'an' needs to be replaced with 'a/an'
    if f" a {target_word}" in context.context.lower():
        context_text = context.context.lower().replace(f" a {target_word}", f" a/an {target_word}")
    elif f" an {target_word}" in context.context.lower():
        context_text = context.context.lower().replace(f" an {target_word}", f" a/an {target_word}")
    else:
        context_text = context.context

    # Bold the first letter of the target word in the gapfill question
    bold_target_letter = f"**{target_word[0]}**{' _' * (len(target_word) - 1)}"
    gapfill_question = context_text.replace(target_word, f"{bold_target_letter} ({translation})")

    # Log the generated gapfill question for debugging
    app.logger.debug(f"Generated gapfill question: {gapfill_question}")

    if request.method == 'POST':
        # Capture the question and answer from the submitted form
        question = request.form['question']
        answer = request.form['answer']

        # Log form data for debugging
        app.logger.debug(f"Form data received - Question: {question}, Answer: {answer}")

        # Save the exercise to the user's list
        exercise = Exercise(
            user_id=session['user_id'],  # Use session's user_id
            question=question,
            answer=answer,
            exercise_type='gapfill'
        )
        db.session.add(exercise)

        # Log the action of adding an exercise
        app.logger.debug(f"Exercise added to the database with question: {question}")

        # Delete the context and reduce the number of contexts
        db.session.delete(context)
        item.number_of_contexts = max(0, item.number_of_contexts - 1)  # Ensure it doesn't go below 0

        # Commit changes to the database
        db.session.commit()

        # Log the deletion of the context and the update of number_of_contexts
        app.logger.debug(f"Context deleted, remaining number of contexts for item {item.item}: {item.number_of_contexts}")

        flash('Exercise saved successfully!', 'success')
        
        # Redirect to the personal space page, passing the redirect_back parameter
        return redirect(url_for('personal_space', username=session['username'], redirect_back=redirect_back))

    # Render the template for creating a gapfill, passing necessary parameters
    return render_template('create_gapfill.html', question=gapfill_question, answer=target_word, redirect_back=redirect_back)





@app.route('/create_mcq/<int:context_id>', methods=['GET', 'POST'])
@csrf.exempt
@login_required
def create_mcq(context_id):
    context = ContextEntry.query.get_or_404(context_id)
    item = VocabularyEntry.query.get_or_404(context.item_id)
    
    # Retrieve the 'redirect_back' parameter, defaulting to the 'generate_exercises' page if not provided
    redirect_back = request.args.get('redirect_back', url_for('generate_exercises', item_id=item.id))
    
    # Log various details for debugging
    app.logger.debug(f"Creating MCQ for context ID: {context_id}")
    app.logger.debug(f"Associated vocabulary item: {item.item}")
    app.logger.debug(f"Context: {context.context}")
    app.logger.debug(f"Redirect back URL: {redirect_back}")
    
    def get_distractors(item, exclude_list, num_distractors=4):
        distractors_query = VocabularyEntry.query.filter(
            VocabularyEntry.course_code == item.course_code,
            VocabularyEntry.pos == item.pos,
            VocabularyEntry.id != item.id,
            VocabularyEntry.selected != 'yes',
            VocabularyEntry.used_as_distractor != 'yes',
            ~VocabularyEntry.item.in_(exclude_list)
        ).all()

        distractors = [d.item for d in distractors_query]
        if len(distractors) < num_distractors:
            distractors += ['-------'] * (num_distractors - len(distractors))

        return random.sample(distractors, min(num_distractors, len(distractors)))
    
    def inflect_word(word, target_word, pos):
        if pos == 'noun':
            if p.plural(target_word) == target_word:
                return p.plural(word)
            else:
                return word
        elif pos == 'verb':
            if target_word.endswith('ed'):
                if word.endswith('e'):
                    return word + 'd'
                elif word.endswith('y') and not word.endswith('ay') and not word.endswith('ey') and not word.endswith('oy'):
                    return word[:-1] + 'ied'
                else:
                    return word + 'ed'
            elif target_word.endswith('s'):
                return word + 's'
            elif target_word.endswith('ing'):
                if word.endswith('e') and len(word) > 1 and word[-2] not in 'aeiou':
                    return word[:-1] + 'ing'
                elif word.endswith('ie'):
                    return word[:-2] + 'ying'
                else:
                    return word + 'ing'
            else:
                return word
        return word
    
    def generate_mcq(context, item):
        if f" a {context.target_word}" in context.context.lower():
            context_text = context.context.lower().replace(f" a {context.target_word}", f" a/an {context.target_word}")
        elif f" an {context.target_word}" in context.context.lower():
            context_text = context.context.lower().replace(f" an {context.target_word}", f" a/an {context.target_word}")
        else:
            context_text = context.context

        distractors = get_distractors(item, [])
        if item.pos in ['noun', 'verb']:
            distractors = [inflect_word(distractor, context.target_word, item.pos) for distractor in distractors]
        
        all_options = [context.target_word] + distractors
        random.shuffle(all_options)
        question = context_text.replace(context.target_word, '__________')
        return question, all_options

    if request.method == 'POST':
        if request.is_json:
            data = request.get_json()
            if data.get('action') == 'change_distractors':
                selected_distractors = data.get('selected_distractors', [])
                current_options = data.get('current_options', [])

                new_distractors = get_distractors(item, [item.item] + [opt for opt in current_options if opt not in selected_distractors], len(selected_distractors))
                if item.pos in ['noun', 'verb']:
                    new_distractors = [inflect_word(distractor, context.target_word, item.pos) for distractor in new_distractors]
                
                updated_options = [opt if opt not in selected_distractors else new_distractors.pop(0) for opt in current_options]
                return jsonify({'new_options': updated_options})

        question = request.form['question']
        answer = request.form['answer']
        options = request.form.getlist('options')
        options_str = '\n'.join([f"{chr(65 + i)}. {option}" for i, option in enumerate(options)])
        
        # Log form data for debugging
        app.logger.debug(f"Form data received - Question: {question}, Answer: {answer}")

        # Marking the used distractors in the database
        for option in options:
            if option != answer and option != '-------':
                distractor_entry = VocabularyEntry.query.filter_by(item=option, course_code=item.course_code).first()
                if distractor_entry:
                    distractor_entry.used_as_distractor = 'yes'

        db.session.commit()

        # Delete the context and reduce the number of contexts
        db.session.delete(context)
        item.number_of_contexts = max(0, item.number_of_contexts - 1)

        db.session.commit()

        # Log the deletion of the context and the update of number_of_contexts
        app.logger.debug(f"Context deleted, remaining number of contexts for item {item.item}: {item.number_of_contexts}")

        exercise = Exercise(
            user_id=session['user_id'],
            question=f"{question}\n{options_str}",
            answer=answer,
            exercise_type='mcq'
        )
        db.session.add(exercise)
        db.session.commit()

        flash(f"Exercise saved successfully!", 'success')
        # Pass the redirect_back parameter to the personal space page
        return redirect(url_for('personal_space', username=session['username'], redirect_back=redirect_back))

    question, options = generate_mcq(context, item)
    answer = context.target_word

    return render_template('create_mcq.html', question=question, options=options, answer=answer, context_id=context_id, redirect_back=redirect_back)





@app.template_filter('to_letter')
def to_letter(value):
    return chr(value + 64)


@app.template_filter('plus_64')
def plus_64(value):
    return chr(value + 64)


@app.route('/generate_exercises/<int:item_id>', methods=['GET', 'POST'])
@csrf.exempt
@login_required
def generate_exercises(item_id):
    # Get the search parameters from the request args to preserve them
    course_code = request.args.getlist('course_code')
    domain = request.args.getlist('domain')
    cefr_level = request.args.getlist('cefr_level')
    number_of_contexts = request.args.getlist('number_of_contexts')
    item_search = request.args.get('item')
    selected = request.args.getlist('selected')

    item = VocabularyEntry.query.get_or_404(item_id)

    # Get current contexts with pagination
    page = request.args.get('page', 1, type=int)
    per_page = 10
    pagination = ContextEntry.query.filter_by(item_id=item_id).paginate(page=page, per_page=per_page, error_out=False)
    current_contexts = pagination.items

    # Find and highlight target words
    highlighted_contexts = []
    for context in current_contexts:
        target_word = find_target_word(context.context, item.item)
        highlighted_context = re.sub(f"\\b{target_word}\\b", f"<b>{target_word}</b>", context.context, flags=re.IGNORECASE)
        highlighted_context_obj = {
            'id': context.id,
            'context': highlighted_context
        }
        highlighted_contexts.append(highlighted_context_obj)

    if request.method == 'POST':
        # Handle exercise generation and saving
        selected_context_id = request.form.get('context_id')
        context = ContextEntry.query.get_or_404(selected_context_id)

        # Save the exercise to the Exercise table
        new_exercise = Exercise(
            user_id=current_user.id,  # Assuming you're using Flask-Login for user management
            question=request.form.get('question'),
            answer=request.form.get('answer'),
            exercise_type=request.form.get('exercise_type'),
        )
        db.session.add(new_exercise)

        # Delete the context after saving the exercise
        db.session.delete(context)

        # Decrement the number_of_contexts in the related vocabulary item
        item.number_of_contexts = max(0, item.number_of_contexts - 1)  # Ensure it doesn't go below 0

        # Commit the changes to the database
        db.session.commit()

        flash('Exercise saved and context deleted.', 'success')
        return redirect(url_for('generate_exercises', item_id=item.id))

    return render_template(
        'generate_exercises.html',
        item=item,
        current_contexts=highlighted_contexts,
        pagination=pagination,
        per_page=per_page,
        page=page,
        course_code=course_code,
        domain=domain,
        cefr_level=cefr_level,
        number_of_contexts=number_of_contexts,
        item_search=item_search,
        selected=selected
    )




@app.route('/my_exercises/<username>')
@login_required
def my_exercises(username):
    if 'username' not in session or session['username'] != username:
        return redirect(url_for('login'))

    user = User.query.filter_by(username=username).first_or_404()
    exercises = Exercise.query.filter_by(user_id=user.id).all()

    for exercise in exercises:
        exercise.question = format_options(exercise.question)

    return render_template('my_exercises.html', exercises=exercises)


@app.route('/edit_exercise', methods=['POST'])
@csrf.exempt
@login_required
def edit_exercise():
    data = request.get_json()
    exercise_id = data.get('exercise_id')
    field = data.get('field')
    new_value = data.get('new_value')

    exercise = Exercise.query.get_or_404(exercise_id)
    if field == 'question':
        exercise.question = new_value
    elif field == 'answer':
        exercise.answer = new_value
    db.session.commit()

    return jsonify({'success': True})


@app.route('/delete_exercises', methods=['POST'])
@csrf.exempt
@login_required
def delete_exercises():
    # Log the complete form data
    app.logger.debug(f"Form data received: {request.form}")
    
    # Extract selected_exercises from the form data
    exercise_ids = request.form.getlist('selected_exercises')
    
    # Log the extracted exercise_ids
    app.logger.debug(f"Extracted exercise_ids after getlist: {exercise_ids}")
    
    if not exercise_ids:
        flash('No exercises have been selected.', 'error')
        return redirect(url_for('my_exercises', username=session['username']))
    
    # Proceed to delete selected exercises
    for exercise_id in exercise_ids:
        app.logger.debug(f"Attempting to delete exercise with ID: {exercise_id}")
        exercise = Exercise.query.get_or_404(exercise_id)
        db.session.delete(exercise)
    
    db.session.commit()

    flash('Selected exercises deleted successfully!', 'success')
    return redirect(url_for('my_exercises', username=session['username']))





@app.route('/export_to_word', methods=['POST'])
@csrf.exempt
@login_required
def export_to_word():
    include_answers = request.form.get('action') == 'export_with_answers'
    selected_exercises = request.form.getlist('selected_exercises')

    # Log form data and extracted selected_exercises
    app.logger.debug(f"Form data: {request.form}")
    app.logger.debug(f"Extracted exercise_ids for export: {selected_exercises}")

    if not selected_exercises:
        flash('No exercises have been selected.', 'error')
        return redirect(url_for('my_exercises', username=session['username']))

    exercises = Exercise.query.filter(Exercise.id.in_(selected_exercises)).all()

    doc = Document()
    doc.add_heading(f"{session['username']}'s Selected Exercises", 0)

    for idx, exercise in enumerate(exercises, 1):
        # Check if the exercise is a gapfill and handle bold formatting
        if exercise.exercise_type == 'gapfill':
            # Use a split method to identify bold text (assuming '**' as the marker)
            parts = re.split(r'(\*\*.*?\*\*)', exercise.question)
            paragraph = doc.add_paragraph(f"{idx}. ")
            
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    # Remove the '**' markers and make the text bold
                    part_text = part.strip("**")
                    paragraph.add_run(part_text).bold = True
                else:
                    # Add regular text
                    paragraph.add_run(part)
        else:
            doc.add_paragraph(f"{idx}. {exercise.question}")
        
        if include_answers:
            doc.add_paragraph(f"Answer: {exercise.answer}")

    if not os.path.exists('exports'):
        os.makedirs('exports')

    file_path = f'exports/{session["username"]}_selected_exercises.docx'
    doc.save(file_path)

    return send_file(file_path, as_attachment=True)








@app.template_filter('escapejs')
def escapejs_filter(s):
    if s is None:
        return ''
    return json.dumps(s)[1:-1]


def format_options(question):
    # Adjust the pattern to handle line breaks and spaces appropriately
    pattern = r'(A\..*?)(?:\s*)(B\..*?)(?:\s*)(C\..*?)(?:\s*)(D\..*?)(?:\s*)(E\..*?)'
    match = re.search(pattern, question, re.DOTALL)
    if match:
        print("Match found:", match.groups())  # Debugging output
        formatted_options = "\t".join(match.groups())  # Join options with a tab
        formatted_question = re.sub(pattern, formatted_options, question, flags=re.DOTALL)
        print("Formatted question:", formatted_question)  # Debugging output
        return formatted_question
    print("No match found for question:", question)  # Debugging output
    return question



if __name__ == '__main__':
    from nlp_utils import update_pos
    if not os.path.exists(app.config['UPLOAD_FOLDER']):
        os.makedirs(app.config['UPLOAD_FOLDER'])
    #update_pos(app)  # UNCOMMENT TO AUTOMATICALLY UPDATE ALL POS IN THE DATABASE
    app.run(debug=True)

