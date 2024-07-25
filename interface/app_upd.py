from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.sql import func
from flask_bcrypt import Bcrypt
from flask_migrate import Migrate
from flask_wtf.csrf import CSRFProtect, CSRFError
import pandas as pd
import os
import re
from datetime import datetime
import nltk
from nltk import pos_tag
from nltk.tokenize import word_tokenize
from nltk.corpus import wordnet
from sqlalchemy import or_
import logging
import json

#in order to export files to Word
from docx import Document
from io import BytesIO
from flask import send_file



nltk.download('averaged_perceptron_tagger')

nltk.download('punkt')
nltk.download('wordnet')

app = Flask(__name__)
csrf = CSRFProtect(app)
app.config['SECRET_KEY'] = 'fj90if0'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///site.db'
app.config['UPLOAD_FOLDER'] = 'uploads'  # Folder to store uploaded files
db = SQLAlchemy(app)
migrate = Migrate(app, db)
bcrypt = Bcrypt(app)

logging.basicConfig(level=logging.DEBUG)


import time
import spacy
import nltk
import google.generativeai as genai
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.stem import PorterStemmer
from nltk.corpus import wordnet
from flask import jsonify
import langid
import random
from difflib import SequenceMatcher
from nltk.stem import WordNetLemmatizer
nltk.download('punkt')
nltk.download('wordnet')

lemmatizer = WordNetLemmatizer()

nlp = spacy.load("en_core_web_sm")
genai.configure(api_key="AIzaSyAFsbEx5fpzsoePD_1Wyct63A8PgBejBKI")  
model = genai.GenerativeModel('gemini-pro')

# Function to get wordnet POS tag
def get_wordnet_pos(tag):
    if tag.startswith('J'):
        return wordnet.ADJ
    elif tag.startswith('V'):
        return wordnet.VERB
    elif tag.startswith('N'):
        return wordnet.NOUN
    elif tag.startswith('R'):
        return wordnet.ADV
    else:
        return wordnet.NOUN  # Default to noun if the POS tag is not recognized


def find_target_word(context, vocabulary_item):
    # Tokenize the context into words
    words = re.findall(r'\w+', context.lower())
    vocabulary_item_lower = vocabulary_item.lower()

    # Check if the vocabulary item is in the context
    if vocabulary_item_lower in words:
        return vocabulary_item

    # Check for lemma match
    for word in words:
        if lemmatizer.lemmatize(word) == vocabulary_item_lower:
            return word

    # Find the word that shares the most letters with the vocabulary item
    most_similar_word = max(words, key=lambda word: SequenceMatcher(None, word, vocabulary_item_lower).ratio())
    return most_similar_word


# Function to lemmatize sentence
def lemmatize_sentence(sentence):
    doc = nlp(sentence)
    lemmatized_words = [token.lemma_ for token in doc]
    lemmatized_sentence = ' '.join(lemmatized_words)
    return lemmatized_sentence

def check_linguistic_features(example, level):
    # Tokenize the example into sentences and words
    sentences = sent_tokenize(example)
    words = [word for word in word_tokenize(example) if word.isalpha()]

    # Initialize a flag to track if all criteria are met
    all_within_range = True

    # 1. Words per sentence
    words_per_sentence = len(words) / len(sentences)
    print(f"Words per sentence: {words_per_sentence}")
    if level == 'B1':
        if (words_per_sentence >= 8) and (words_per_sentence <= 26.05): #lower limit needed otherwise too short; 1.5 std used and here not 'per sentence' but for example in reference corpus
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False
    if level == 'B2':
        if (words_per_sentence >= 17) and (words_per_sentence <= 33.69):
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False

    # 2. Letters per word
    letters_per_word = sum(len(word) for word in words) / len(words)
    print(f"Letters per word: {letters_per_word}")
    if level == 'B1':
        if letters_per_word <= 9.26:
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False
    if level == 'B2':
        if letters_per_word <= 10.96:
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False

    # 3. Noun phrases per sentence
    total_noun_phrases = sum(len(list(nlp(sent).noun_chunks)) for sent in sentences)
    noun_phrases_per_sentence = total_noun_phrases / len(sentences)
    print(f"Noun phrases per sentence: {noun_phrases_per_sentence}")
    if level == 'B1':
        if noun_phrases_per_sentence <= 10.49:
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False
    if level == 'B2':
        if (noun_phrases_per_sentence >= 0.47) and (noun_phrases_per_sentence <= 11.71):
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False

    # 4. Non-stem words per sentence
    stemmer = PorterStemmer()
    total_non_stem_words = sum(1 for word in words if word.lower() != stemmer.stem(word.lower()))
    non_stem_words_per_sentence = total_non_stem_words / len(sentences)
    print(f"Non-stem words per sentence: {non_stem_words_per_sentence}")
    if level == 'B1':
        if (non_stem_words_per_sentence >= 2.47) and (non_stem_words_per_sentence <= 55.52): #lower limit is problematic! using 3 std for it; both B1 and B2, which brings B2 to 0
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False
    if level == 'B2':
        if non_stem_words_per_sentence <= 54.46:
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False

    # 5. Punctuation signs per sentence
    punctuation_to_count = "',:;-\"‘’“”"
    total_punctuation = sum(example.count(char) for char in punctuation_to_count)
    punctuation_per_sentence = total_punctuation / len(sentences)
    print(f"Punctuation signs per sentence: {punctuation_per_sentence}")
    if level == 'B1':
        if punctuation_per_sentence <= 2.41:
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False
    if level == 'B2':
        if punctuation_per_sentence <= 2.82:
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False

    # 6. Verbs per sentence
    total_verbs = sum(1 for token in nlp(example) if token.pos_ == 'VERB')
    verbs_per_sentence = total_verbs / len(sentences)
    print(f"Verbs per sentence: {verbs_per_sentence}")
    if level == 'B1':
        if verbs_per_sentence <= 4.95:
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False
    if level == 'B2':
        if (verbs_per_sentence >= 0.45) and (verbs_per_sentence <= 4.77):
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False

    # 7. Adjectives and adverbs per sentence
    total_adj_adv = sum(1 for token in nlp(example) if token.pos_ in ('ADJ', 'ADV'))
    adj_adv_per_sentence = total_adj_adv / len(sentences)
    print(f"Adjectives and adverbs per sentence: {adj_adv_per_sentence}")
    if level == 'B1':
        if adj_adv_per_sentence <= 6.4:
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False
    if level == 'B2':
        if (adj_adv_per_sentence >= 0.99) and (adj_adv_per_sentence <= 5.15):
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False

    # 8. First person pronouns per sentence
    first_person_pronouns = {"i", "me", "my", "myself", "we", "us", "our", "ours", "ourselves"}
    total_first_person_pronouns = sum(1 for word in words if word.lower() in first_person_pronouns)
    first_person_pronouns_per_sentence = total_first_person_pronouns / len(sentences)
    print(f"First person pronouns per sentence: {first_person_pronouns_per_sentence}")
    if level == 'B1':
        if first_person_pronouns_per_sentence <= 1:
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False
    if level == 'B2':
        if first_person_pronouns_per_sentence <= 0.79:
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False

    # 9. Proper nouns per sentence
    total_proper_nouns = sum(1 for token in nlp(example) if token.pos_ == 'PROPN')
    proper_nouns_per_sentence = total_proper_nouns / len(sentences)
    print(f"Proper nouns per sentence: {proper_nouns_per_sentence}")
    if level == 'B1':
        if proper_nouns_per_sentence <= 3.39:
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False
    if level == 'B2':
        if proper_nouns_per_sentence <= 5.19:
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False

    # 10. Pronouns per sentence
    total_pronouns = sum(1 for token in nlp(example) if token.pos_ == 'PRON')
    pronouns_per_sentence = total_pronouns / len(sentences)
    print(f"Pronouns per sentence: {pronouns_per_sentence}")
    if level == 'B1':
        if pronouns_per_sentence <= 2.65:
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False
    if level == 'B2':
        if pronouns_per_sentence <= 3.11:
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False

    # 11. Anaphora words per sentence
    anaphora_words = set(["the", "he", "she", "it", "they", "this", "that", "these", "those",
                          "who", "which", "whose", "whom", "where", "all",
                          "some", "none", "any", "each", "every", "here", "there", "now", "then"])
    total_anaphora_words = sum(1 for word in word_tokenize(example) if word.lower() in anaphora_words)
    anaphora_words_per_sentence = (total_anaphora_words / len(words)) * 100
    print(f"Anaphora words per sentence: {anaphora_words_per_sentence}")
    if level == 'B1':
        if (anaphora_words_per_sentence >= 0.13) and (anaphora_words_per_sentence <= 25.77):
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False
    if level == 'B2':
        if anaphora_words_per_sentence <= 23.7:
            print("Falls within the range")
        else:
            print("Doesn't fall within the range")
            all_within_range = False

    # Set example to empty string if not all criteria are met
    if not all_within_range:
        example = ""

    return example

def clean_text(sentences):
    cleaned_sentences = []
    for text in sentences:
        text = text.replace(' ,', ',')
        text = text.replace(' .', '.')
        text = text.replace('\n', ' ')
        text = text.lstrip('#')   
        text = text.lstrip('-')
        text = text.lstrip('•') 
        text = text.lstrip('>')
        text = text.replace('*', '')
        
        # If text starts with up to 3 words followed by ':', remove this part:
        parts = text.split(':', 1)
        if len(parts) > 1 and len(parts[0].split()) <= 3:
            text = parts[1].strip()
            
        # If the sentence now ends in ':.' remove this pattern as well as anything before it following the last end-of-sentence punctuation
        if text.endswith(':.'):
            text = text.replace(':.', '')
            if ('.' in text) or ('!' in text) or ('?' in text):
                last_punctuation_index = max(text.rfind('.'), text.rfind('?'), text.rfind('!'))
                text = text[:last_punctuation_index].strip()

        # Remove text within square brackets at the beginning of the text
        text = re.sub(r'^\[[^\]]*\]\s*', '', text)
        
        if 'Fig.' in text:
            open_bracket_index = text.find('(')
            if open_bracket_index != -1 and text.count('(') > text.count(')'):
                text = text[:open_bracket_index].strip()

        # Remove opening quotation sign if not closed
        if text.startswith(("'", '"', '“', '‘', '„', '‹', '«')):
            if "'" not in text[1:] and '"' not in text[1:]:
                text = text[1:]

        # Capitalize the first letter if it's not already capitalized
        if text and not text[0].isupper():
            text = text[0].upper() + text[1:]
        
        # Remove short text in beginning like enumeration
        match_short_text = re.match(r'^[a-zA-Z0-9\s]{1,10}\.\s*', text)
        if match_short_text:
            text = text[len(match_short_text.group()):]

        # Put a full stop if no punctuation at the end 
        if not text.endswith(('.', '!', '?')):
            text += '.'

        # Strip extra spaces
        text = text.strip()

        cleaned_sentences.append(text)
    return cleaned_sentences



def generate_and_validate_contexts(item, pos, domain, level):
    contexts = []
    lower = 8 if level == 'B1' else 17
    upper = 43 if level == 'B1' else 87

    question_template = "Here is a sentence at level {level} showing how you use the {pos} '{item}' in the domain of {domain} ({lower}-{upper} words):"
    question = question_template.format(
        level=level,
        pos=pos,
        item=item,
        domain=domain,
        lower=lower,
        upper=upper
    )

    while len(contexts) < 5:  # Generate 5 valid contexts
        response = model.generate_content(question)
        context = response.text if response else ''
        context = context.strip()

        if not context:
            continue

        # Perform various checks
        if pos in ["noun", "verb", "expression"]:
            context_lemmas = lemmatize_sentence(context).lower()
        else:
            context_lemmas = context.lower()

        if item not in context_lemmas and lemmatize_sentence(item) not in context_lemmas:
            continue

        context = re.sub(r'^\d+(\.|\))', '', context)
        context = re.sub(r'^[^\w\s]+(?=[A-Z])', '', context)
        context = re.sub(r'Here is a sentence.*?:', '', context)
        context = re.sub(r'^[\'\"“‘]+|[\'\"”’]+$', '', context)
        context = re.sub(r'\.[\'"”’]+\.$', '.', context)

        sentences = sent_tokenize(context)
        if len(sentences) > 1 and sentences[-1].startswith(('In this sentence', 'In the sentence')):
            sentences.pop()
        context = ' '.join(sentences)

        patterns = [f"{item} is", f"{item} refers to", f"A {item.lower()} is", f"A {item.lower()} refers to",
                    f"An {item.lower()} is", f"An {item.lower()} refers to"]
        if any(context.lower().startswith(pattern) for pattern in patterns):
            continue

        if any(keyword in context for keyword in ['Example:', 'Answer:', 'CEFR', 'A1', 'A2', 'B1', 'B2', 'C1', 'C2']):
            continue

        tokens = word_tokenize(context)
        pos_tags = pos_tag(tokens)

        if pos in ["verb", "noun", "adjective", "adverb"]:
            if not any((word.lower() == item.lower() and
                        ((pos == 'noun' and tag.startswith('NN')) or
                         (pos == 'verb' and tag.startswith('VB')) or
                         (pos == 'adjective' and tag.startswith('JJ')) or
                         (pos == 'adverb' and tag.startswith('RB'))))
                       for word, tag in pos_tags):
                continue

        threshold = len(context) // 4
        non_latin_count = sum(1 for char in context if not char.isascii())
        if non_latin_count > threshold:
            continue

        lang, _ = langid.classify(context)
        if lang != 'en':
            continue

        if context.count(':') > 1:
            continue

        sentence = context.lower()
        base_form = item.lower()
        base_lemma = lemmatize_sentence(base_form)
        words = word_tokenize(sentence)
        lemmatized_words = [lemmatize_sentence(word) for word in words]
        if words.count(base_form) > 1 or lemmatized_words.count(base_lemma) > 1:
            continue

        if context.startswith('*') or context.startswith('-'):
            context = context[1:]

        context = context.strip()
        if not context.endswith(('.', '?', '!')):
            context += '.'

        context = check_linguistic_features(context, level)
        if context:
            contexts.append(context)

    return clean_text(contexts)



class User(db.Model):
    __tablename__ = 'users'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(20), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(128), nullable=False)

    def set_password(self, password):
        self.password_hash = bcrypt.generate_password_hash(password).decode('utf-8')

    def check_password(self, password):
        return bcrypt.check_password_hash(self.password_hash, password)


class VocabularyEntry(db.Model):
    __tablename__ = 'vocabulary_entries'
    id = db.Column(db.Integer, primary_key=True)
    item = db.Column(db.String(255), nullable=False)
    pos = db.Column(db.String(50), nullable=True)
    translation = db.Column(db.String(255), nullable=True)
    lesson_title = db.Column(db.String(255), nullable=True)
    reading_or_listening = db.Column(db.String(50), nullable=True)
    course_code = db.Column(db.String(50), nullable=False)
    cefr_level = db.Column(db.String(10), nullable=False)
    domain = db.Column(db.String(255), nullable=False)
    user = db.Column(db.String(50), nullable=False)
    date_loaded = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)
    number_of_contexts = db.Column(db.Integer, nullable=False, default=0)
    contexts = db.relationship('ContextEntry', backref='vocabulary_entry', lazy=True)


class ContextEntry(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    item_id = db.Column(db.Integer, db.ForeignKey('vocabulary_entries.id'), nullable=False)
    context = db.Column(db.Text, nullable=False)
    target_word = db.Column(db.String(128), nullable=False)
    user = db.Column(db.String(50), nullable=False)
    date_added = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)



class Exercise(db.Model):
    __tablename__ = 'exercises'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'), nullable=False)
    question = db.Column(db.Text, nullable=False)
    answer = db.Column(db.String(128), nullable=False)
    exercise_type = db.Column(db.String(50), nullable=False)  # e.g., gapfill, mcq
    date_created = db.Column(db.DateTime, default=datetime.utcnow)

    user = db.relationship('User', backref=db.backref('exercises', lazy=True))



def strip_prefixes(item):
    prefixes = ['to ', 'a ', 'an ', 'the ']
    for prefix in prefixes:
        if item.lower().startswith(prefix):
            return item[len(prefix):]
    return item

def get_wordnet_pos(treebank_tag):
    if treebank_tag.startswith('J'):
        return 'adjective'
    elif treebank_tag.startswith('V'):
        return 'verb'
    elif treebank_tag.startswith('N'):
        return 'noun'
    elif treebank_tag.startswith('R'):
        return 'adverb'
    else:
        return 'word'

def update_pos():
    with app.app_context():
        entries = VocabularyEntry.query.all()
        for entry in entries:
            if not entry.pos or entry.pos == 'word':
                if ' ' in entry.item:
                    entry.pos = 'expression'
                else:
                    nltk_tag = pos_tag([entry.item])[0][1]
                    entry.pos = get_wordnet_pos(nltk_tag)
        db.session.commit()

@app.route('/')
def home():
    return render_template('home.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        email = request.form['email']
        password = request.form['password']

        existing_user = User.query.filter((User.username == username) | (User.email == email)).first()
        if existing_user:
            flash("Username or email already registered.", 'error')
            return render_template('register.html')

        new_user = User(username=username, email=email)
        new_user.set_password(password)
        db.session.add(new_user)
        db.session.commit()
        flash('Registration successful. Please log in.', 'success')
        return redirect(url_for('login'))

    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if 'username' in session:
        return redirect(url_for('home'))

    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')

        user = User.query.filter_by(username=username).first()
        if user and user.check_password(password):
            session['username'] = username
            session['user_id'] = user.id
            return redirect(url_for('home'))
        else:
            flash('Invalid username or password', 'error')

    return render_template('login.html')


@app.route('/logout', methods=['GET', 'POST'])
def logout():
    session.clear()
    return redirect(url_for('logout_confirmation'))

@app.route('/logout_confirmation')
def logout_confirmation():
    return """
    <p>You have been logged out successfully.</p>
    <p><a href="/">Back to Home Page</a></p>
    """

@app.route('/personal_space/<username>')
def personal_space(username):
    if 'username' not in session or session['username'] != username:
        return redirect(url_for('login'))
    return render_template('personal_space.html', username=username)


@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/add_vocabulary_list', methods=['GET', 'POST'])
def add_vocabulary_list():
    if 'username' not in session:
        return redirect(url_for('login'))

    if request.method == 'POST':
        course_code = request.form['course_code']
        cefr_level = request.form['cefr_level']
        domain = request.form['domain']
        file = request.files['file']

        if file:
            filename = file.filename
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)

            if filename.endswith('.xlsx'):
                df = pd.read_excel(filepath)
            elif filename.endswith('.csv'):
                df = pd.read_csv(filepath)
            else:
                flash('Invalid file format. Please upload a .csv or .xlsx file.', 'error')
                return redirect(url_for('add_vocabulary_list'))

            required_columns = ['Item', 'POS', 'Translation', 'Lesson title', 'Reading or listening']
            for column in required_columns:
                if column not in df.columns:
                    df[column] = ''

            user = session['username']
            date_loaded = datetime.utcnow()
            number_of_contexts = 0
            contexts = ''

            for index, row in df.iterrows():
                entry = VocabularyEntry(
                    item=strip_prefixes(row['Item']),
                    pos=row['POS'] if row['POS'] else None,
                    translation=row['Translation'],
                    lesson_title=row['Lesson title'],
                    reading_or_listening=row['Reading or listening'],
                    course_code=course_code,
                    cefr_level=cefr_level,
                    domain=domain,
                    user=user,
                    date_loaded=date_loaded,
                    number_of_contexts=number_of_contexts,
                    contexts=contexts
                )
                db.session.add(entry)
            db.session.commit()
            flash('Thank you for uploading a vocabulary list!', 'success')
            return redirect(url_for('home'))

    return render_template('add_vocabulary_list.html')


@app.route('/add_context_examples', methods=['GET', 'POST'])
@csrf.exempt  # Disable CSRF protection for this route
def add_context_examples():
    if 'username' not in session:
        return redirect(url_for('login'))

    # Get unique values for dropdowns and sort them alphabetically
    course_codes = sorted([code.course_code for code in VocabularyEntry.query.with_entities(VocabularyEntry.course_code).distinct().all()])
    domains = sorted([domain.domain for domain in VocabularyEntry.query.with_entities(VocabularyEntry.domain).distinct().all()])
    levels = sorted([level.cefr_level for level in VocabularyEntry.query.with_entities(VocabularyEntry.cefr_level).distinct().all()])
    items = sorted([strip_prefixes(item.item) for item in VocabularyEntry.query.with_entities(VocabularyEntry.item).distinct().all()])

    results = []
    selected_params = {
        'course_code': [],
        'domain': [],
        'cefr_level': [],
        'number_of_contexts': [],
        'item': ''
    }

    if request.method == 'POST':
        search_params = {
            'course_code': request.form.getlist('course_code'),
            'domain': request.form.getlist('domain'),
            'cefr_level': request.form.getlist('cefr_level'),
            'item': request.form.get('item'),  # Get the single selected item
            'number_of_contexts': request.form.getlist('number_of_contexts')
        }
        selected_params = search_params

        query = VocabularyEntry.query
        if search_params['course_code']:
            query = query.filter(VocabularyEntry.course_code.in_(search_params['course_code']))
        if search_params['domain']:
            query = query.filter(VocabularyEntry.domain.in_(search_params['domain']))
        if search_params['cefr_level']:
            query = query.filter(VocabularyEntry.cefr_level.in_(search_params['cefr_level']))
        if search_params['item']:
            query = query.filter(VocabularyEntry.item == search_params['item'])
        if search_params['number_of_contexts']:
            context_filters = []
            for context_range in search_params['number_of_contexts']:
                if context_range == 'less_than_10':
                    context_filters.append(VocabularyEntry.number_of_contexts < 10)
                elif context_range == '10_to_20':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(10, 20))
                elif context_range == '20_to_50':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(20, 50))
                elif context_range == '50_to_100':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(50, 100))
            query = query.filter(or_(*context_filters))

        results = query.order_by(VocabularyEntry.item).all()  # Sort results alphabetically by item

        # Add actual count of context examples to each result
        for result in results:
            context_count = ContextEntry.query.filter_by(item_id=result.id).count()
            result.number_of_contexts = context_count

    return render_template('add_context_examples.html',
                           course_codes=course_codes,
                           domains=domains,
                           levels=levels,
                           items=items,
                           results=results,
                           selected_params=selected_params)



@app.route('/current_contexts/<int:item_id>', methods=['GET'])
def current_contexts(item_id):
    item = VocabularyEntry.query.get_or_404(item_id)
    contexts = ContextEntry.query.filter_by(item_id=item_id).all()
    return render_template('current_contexts.html', item=item, contexts=contexts)



@app.route('/generate_context/<int:item_id>', methods=['GET', 'POST'])
@csrf.exempt
def generate_context(item_id):
    item = VocabularyEntry.query.get_or_404(item_id)
    generated_contexts = None

    if request.method == 'POST':
        if 'generate_contexts' in request.form:
            pos = item.pos
            domain = item.domain
            level = item.cefr_level
            generated_contexts = generate_and_validate_contexts(item.item, pos, domain, level)

        elif 'save_contexts' in request.form:
            selected_contexts = request.form.getlist('selected_contexts')
            context_texts = request.form.getlist('context_texts')
            user = session.get('username')
            date_added = datetime.utcnow()
            for idx in selected_contexts:
                idx = int(idx) - 1
                if idx < len(context_texts):
                    context = context_texts[idx]
                    target_word = find_target_word(context, item.item)
                    new_context = ContextEntry(item_id=item_id, context=context, target_word=target_word, user=user, date_added=date_added)
                    db.session.add(new_context)
            db.session.commit()
            flash('Selected contexts saved successfully.', 'success')

    # Get current contexts with pagination
    page = request.args.get('page', 1, type=int)
    per_page = 10
    pagination = ContextEntry.query.filter_by(item_id=item_id).paginate(page=page, per_page=per_page, error_out=False)
    current_contexts = pagination.items

    # Find and highlight target words
    highlighted_contexts = []
    for context in current_contexts:
        target_word = find_target_word(context.context, item.item)
        highlighted_context = re.sub(f"\\b{target_word}\\b", f"<b>{target_word}</b>", context.context, flags=re.IGNORECASE)
        highlighted_contexts.append(highlighted_context)

    return render_template('generate_context.html', item=item, current_contexts=highlighted_contexts, generated_contexts=generated_contexts, pagination=pagination, per_page=per_page, page=page, find_target_word=find_target_word)



@app.route('/create_exercises', methods=['GET', 'POST'])
@csrf.exempt
def create_exercises():
    if 'username' not in session:
        return redirect(url_for('login'))

    # Get unique values for checkboxes and sort them alphabetically
    course_codes = sorted([code.course_code for code in VocabularyEntry.query.with_entities(VocabularyEntry.course_code).distinct().all()])
    domains = sorted([domain.domain for domain in VocabularyEntry.query.with_entities(VocabularyEntry.domain).distinct().all()])
    levels = sorted([level.cefr_level for level in VocabularyEntry.query.with_entities(VocabularyEntry.cefr_level).distinct().all()])
    context_ranges = ['less_than_10', '10_to_20', '20_to_50', '50_to_100']

    items = sorted([strip_prefixes(item.item) for item in VocabularyEntry.query.with_entities(VocabularyEntry.item).distinct().all()])

    results = []
    selected_params = {}

    if request.method == 'POST':
        search_params = {
            'course_code': request.form.getlist('course_code'),
            'domain': request.form.getlist('domain'),
            'cefr_level': request.form.getlist('cefr_level'),
            'item': request.form.get('item'),
            'number_of_contexts': request.form.getlist('number_of_contexts')
        }
        selected_params = search_params

        query = VocabularyEntry.query
        if search_params['course_code']:
            query = query.filter(VocabularyEntry.course_code.in_(search_params['course_code']))
        if search_params['domain']:
            query = query.filter(VocabularyEntry.domain.in_(search_params['domain']))
        if search_params['cefr_level']:
            query = query.filter(VocabularyEntry.cefr_level.in_(search_params['cefr_level']))
        if search_params['item']:
            query = query.filter(VocabularyEntry.item == search_params['item'])
        if search_params['number_of_contexts']:
            context_filters = []
            for context_range in search_params['number_of_contexts']:
                if context_range == 'less_than_10':
                    context_filters.append(VocabularyEntry.number_of_contexts < 10)
                elif context_range == '10_to_20':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(10, 20))
                elif context_range == '20_to_50':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(20, 50))
                elif context_range == '50_to_100':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(50, 100))
            query = query.filter(or_(*context_filters))

        results = query.order_by(VocabularyEntry.item).all()  # Sort results alphabetically by item

    return render_template('create_exercises.html',
                           course_codes=course_codes,
                           domains=domains,
                           levels=levels,
                           context_ranges=context_ranges,
                           items=items,
                           results=results,
                           selected_params=selected_params)



@app.route('/filter_items', methods=['POST'])
def filter_items():
    try:
        data = request.get_json()
        logging.debug(f"Received data: {data}")
        course_codes = data.get('course_code', [])
        domains = data.get('domain', [])
        cefr_levels = data.get('cefr_level', [])
        context_ranges = data.get('number_of_contexts', [])

        query = VocabularyEntry.query

        if course_codes:
            query = query.filter(VocabularyEntry.course_code.in_(course_codes))
        if domains:
            query = query.filter(VocabularyEntry.domain.in_(domains))
        if cefr_levels:
            query = query.filter(VocabularyEntry.cefr_level.in_(cefr_levels))
        if context_ranges:
            context_filters = []
            for context_range in context_ranges:
                if context_range == 'less_than_10':
                    context_filters.append(VocabularyEntry.number_of_contexts < 10)
                elif context_range == '10_to_20':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(10, 20))
                elif context_range == '20_to_50':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(20, 50))
                elif context_range == '50_to_100':
                    context_filters.append(VocabularyEntry.number_of_contexts.between(50, 100))
            query = query.filter(or_(*context_filters))

        items = sorted([strip_prefixes(item.item) for item in query.with_entities(VocabularyEntry.item).distinct().all()])

        logging.debug(f"Filtered items: {items}")
        return jsonify(items)
    except Exception as e:
        logging.error(f"Error in filter_items: {e}")
        return jsonify({'error': str(e)}), 400





@app.route('/create_gapfill/<int:context_id>', methods=['GET', 'POST'])
@csrf.exempt  # Disable CSRF protection for this route
def create_gapfill(context_id):
    context = ContextEntry.query.get_or_404(context_id)
    item = VocabularyEntry.query.get_or_404(context.item_id)
    
    target_word = context.target_word
    translation = item.translation
    gapfill_question = context.context.replace(target_word, f"{target_word[0]}{' _' * (len(target_word) - 1)} ({translation})")

    if request.method == 'POST':
        question = request.form['question']
        answer = request.form['answer']
        
        # Save the exercise to the user's list
        exercise = Exercise(
            user_id=session['user_id'],
            question=question,
            answer=answer,
            exercise_type='gapfill'
        )
        db.session.add(exercise)
        db.session.commit()

        flash('Exercise saved successfully!', 'success')
        return redirect(url_for('personal_space', username=session['username']))

    return render_template('create_gapfill.html', question=gapfill_question, answer=target_word)



@app.route('/create_mcq/<int:context_id>', methods=['GET', 'POST'])
@csrf.exempt
def create_mcq(context_id):
    context = ContextEntry.query.get_or_404(context_id)
    item = VocabularyEntry.query.get_or_404(context.item_id)

    def get_distractors(item, exclude_list, num_distractors=4):
        distractors_query = VocabularyEntry.query.filter(
            VocabularyEntry.course_code == item.course_code,
            VocabularyEntry.pos == item.pos,
            VocabularyEntry.id != item.id,
            ~VocabularyEntry.item.in_(exclude_list)
        ).all()
        
        distractors = [d.item for d in distractors_query]
        if len(distractors) < num_distractors:
            distractors += ['-------'] * (num_distractors - len(distractors))
        
        return random.sample(distractors, min(num_distractors, len(distractors)))

    def generate_mcq(context, item):
        distractors = get_distractors(item, [])
        all_options = [item.item] + distractors
        random.shuffle(all_options)
        question = context.context.replace(context.target_word, '__________')
        return question, all_options

    if request.method == 'POST':
        if request.is_json:
            data = request.get_json()
            if data.get('action') == 'change_distractors':
                selected_distractors = data.get('selected_distractors', [])
                current_options = data.get('current_options', [])

                new_distractors = get_distractors(item, [item.item] + [opt for opt in current_options if opt not in selected_distractors], len(selected_distractors))

                updated_options = [opt if opt not in selected_distractors else new_distractors.pop(0) for opt in current_options]

                return jsonify({'new_options': updated_options})

        question = request.form['question']
        answer = request.form['answer']
        options = request.form.getlist('options')
        options_str = '\n'.join([f"{chr(65 + i)}. {option}" for i, option in enumerate(options)])

        exercise = Exercise(
            user_id=session['user_id'],
            question=f"{question}\n{options_str}",
            answer=answer,
            exercise_type='mcq'
        )
        db.session.add(exercise)
        db.session.commit()

        flash(f"Exercise saved successfully! {session['username']}'s Personal Page", 'success')
        return redirect(url_for('personal_space', username=session['username']))

    question, options = generate_mcq(context, item)
    answer = context.target_word

    return render_template('create_mcq.html', question=question, options=options, answer=answer, context_id=context_id)





@app.template_filter('to_letter')
def to_letter(value):
    return chr(value + 64)


@app.template_filter('plus_64')
def plus_64(value):
    return chr(value + 64)



@app.route('/generate_exercises/<int:item_id>', methods=['GET', 'POST'])
@csrf.exempt
def generate_exercises(item_id):
    item = VocabularyEntry.query.get_or_404(item_id)

    # Get current contexts with pagination
    page = request.args.get('page', 1, type=int)
    per_page = 10
    pagination = ContextEntry.query.filter_by(item_id=item_id).paginate(page=page, per_page=per_page, error_out=False)
    current_contexts = pagination.items

    # Find and highlight target words
    highlighted_contexts = []
    for context in current_contexts:
        target_word = find_target_word(context.context, item.item)
        highlighted_context = re.sub(f"\\b{target_word}\\b", f"<b>{target_word}</b>", context.context, flags=re.IGNORECASE)
        # Create a new object to store the highlighted context and the context ID
        highlighted_context_obj = {
            'id': context.id,
            'context': highlighted_context
        }
        highlighted_contexts.append(highlighted_context_obj)

    return render_template('generate_exercises.html', item=item, current_contexts=highlighted_contexts, pagination=pagination, per_page=per_page, page=page)


@app.route('/my_exercises/<username>')
def my_exercises(username):
    if 'username' not in session or session['username'] != username:
        return redirect(url_for('login'))

    user = User.query.filter_by(username=username).first_or_404()
    exercises = Exercise.query.filter_by(user_id=user.id).all()

    for exercise in exercises:
        exercise.question = format_options(exercise.question)

    return render_template('my_exercises.html', exercises=exercises)


@app.route('/edit_exercise', methods=['POST'])
@csrf.exempt
def edit_exercise():
    data = request.get_json()
    exercise_id = data.get('exercise_id')
    field = data.get('field')
    new_value = data.get('new_value')

    exercise = Exercise.query.get_or_404(exercise_id)
    if field == 'question':
        exercise.question = new_value
    elif field == 'answer':
        exercise.answer = new_value
    db.session.commit()

    return jsonify({'success': True})

@app.route('/delete_exercises', methods=['POST'])
@csrf.exempt
def delete_exercises():
    selected_exercises = request.form.getlist('selected_exercises')
    for exercise_id in selected_exercises:
        exercise = Exercise.query.get_or_404(exercise_id)
        db.session.delete(exercise)
    db.session.commit()

    flash('Selected exercises deleted successfully!', 'success')
    return redirect(url_for('my_exercises', username=session['username']))


@app.route('/export_to_word', methods=['GET', 'POST'])
@csrf.exempt
def export_to_word():
    include_answers = request.form.get('include_answers') == 'true'
    user_id = session['user_id']
    user = User.query.get_or_404(user_id)
    exercises = Exercise.query.filter_by(user_id=user.id).all()

    doc = Document()
    doc.add_heading(f"{user.username}'s Exercises", 0)

    for idx, exercise in enumerate(exercises, 1):
        doc.add_paragraph(f"{idx}. {exercise.question}")
        if include_answers:
            doc.add_paragraph(f"Answer: {exercise.answer}")
        #doc.add_paragraph("\n")

    if not os.path.exists('exports'):
        os.makedirs('exports')

    file_path = f'exports/{user.username}_exercises.docx'
    doc.save(file_path)

    return send_file(file_path, as_attachment=True)





@app.template_filter('escapejs')
def escapejs_filter(s):
    if s is None:
        return ''
    return json.dumps(s)[1:-1]


def format_options(question):
    # Adjust the pattern to handle line breaks and spaces appropriately
    pattern = r'(A\..*?)(?:\s*)(B\..*?)(?:\s*)(C\..*?)(?:\s*)(D\..*?)(?:\s*)(E\..*?)'
    match = re.search(pattern, question, re.DOTALL)
    if match:
        print("Match found:", match.groups())  # Debugging output
        formatted_options = "\t".join(match.groups())  # Join options with a tab
        formatted_question = re.sub(pattern, formatted_options, question, flags=re.DOTALL)
        print("Formatted question:", formatted_question)  # Debugging output
        return formatted_question
    print("No match found for question:", question)  # Debugging output
    return question



if __name__ == '__main__':
    if not os.path.exists(app.config['UPLOAD_FOLDER']):
        os.makedirs(app.config['UPLOAD_FOLDER'])
    with app.app_context(): #UNCOMMENT TO AUTOMATICALLY UPDATE ALL POS IN THE DATABASE
        update_pos()  # update POS
    app.run(debug=True)
 
